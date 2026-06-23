"""Generate bank_forecast_accuracy.docx from data."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return t

# ── Title ──
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MIP — Bank-Level Forecast Accuracy Summary")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Prepared by: Eshan Adyanthaya, MPi  |  23 June 2026")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

# ── 1. Methodology ──
heading("1. Methodology")
method_rows = [
    ("Models", "Facebook Prophet (logistic growth) for banks with structural breaks;\nHolt-Winters ETS for clean, stable-growth banks"),
    ("Validation", "Walk-forward cross-validation — 36-month initial window,\n6-month forecast horizon, 6-month step"),
    ("Bank Coverage", "10 CC banks + 15 DC banks modelled individually (ground-up approach)"),
    ("Industry Coverage", "~91% of CC outstanding, ~83% of DC outstanding"),
    ("Residual", "Remaining banks captured via a single residual Prophet model;\nsum of individual + residual = India total"),
    ("Confidence Intervals", "90% prediction intervals from Prophet / ETS simulation"),
]
add_table(["Item", "Detail"], method_rows)
doc.add_paragraph()

# ── 2. CC ──
heading("2. Credit Card Outstanding — Bank-Level Accuracy")
doc.add_paragraph("10 banks modelled individually, covering ~91% of India CC outstanding.")

# Sources:
# - CV summary CSV: HDFC (5.68), SBI (4.84), Kotak (21.87), BoB (20.81), HSBC (13.40)
# - ETS_BANKS comments: ICICI (2.46), Axis (5.97), IndusInd (4.36)
# - Yes Bank: starts Jun 2020 (~68 months), <72m threshold, CV not run
# - Canara Bank: starts Apr 2020 (~70 months), <72m threshold, CV not run
cc_rows = [
    ("ICICI Bank", "ETS", "2.5%", "Stable growth, ETS outperforms Prophet"),
    ("IndusInd Bank", "ETS", "4.4%", "Stable growth"),
    ("State Bank of India", "Prophet", "4.8%", "38 CV windows, range 3.2%-11.3%"),
    ("HDFC Bank", "Prophet", "5.7%", "38 CV windows, range 3.1%-8.1%"),
    ("Axis Bank", "ETS", "6.0%", "Stable growth"),
    ("HSBC", "Prophet", "13.4%", "65 CV windows, range 6.7%-23.4%"),
    ("Bank of Baroda", "Prophet", "20.8%", "27 CV windows, logistic cap applied"),
    ("Kotak Mahindra Bank", "Prophet", "21.9%", "33 CV windows, logistic cap applied"),
    ("Yes Bank", "Prophet", "N/A*", "Post-moratorium (Jun 2020), insufficient history for CV"),
    ("Canara Bank", "Prophet", "N/A*", "Post-merger (Apr 2020), insufficient history for CV"),
]
add_table(["Bank", "Model", "CV MAPE (Median)", "Notes"], cc_rows)
doc.add_paragraph("* Yes Bank and Canara Bank have <72 months of post-event training data; CV requires 36-month initial window + 6-month horizon. Their forecasts are generated but not yet cross-validated. CV will become available as more months accumulate.")
p = doc.add_paragraph()
run = p.add_run("CC Median (validated banks): 5.7%")
run.bold = True
doc.add_paragraph()

# ── 3. DC ──
heading("3. Debit Card Outstanding — Bank-Level Accuracy")
doc.add_paragraph("15 banks modelled individually, covering ~83% of India DC outstanding.")

# Sources:
# - CV summary: SBI (4.26), BoB (5.28), BoI (9.61), Kotak (7.81), CBI (7.30), Paytm (21.39)
# - ETS_BANKS comments: HDFC (1.16), Axis (3.05), ICICI (5.28), UCO (2.64), IOB (5.70)
# - Short-series (start Apr 2020, ~70 months): Canara, Union Bank, PNB, Indian Bank
dc_rows = [
    ("HDFC Bank", "ETS", "1.2%", "Very stable"),
    ("UCO Bank", "ETS", "2.6%", "Stable"),
    ("Axis Bank", "ETS", "3.1%", "Stable"),
    ("State Bank of India", "Prophet", "4.3%", "38 CV windows, range 1.5%-7.1%"),
    ("Bank of Baroda", "Prophet", "5.3%", "27 CV windows, range 1.5%-7.4%"),
    ("ICICI Bank", "ETS", "5.3%", ""),
    ("Indian Overseas Bank", "ETS", "5.7%", ""),
    ("Central Bank of India", "Prophet", "7.3%", "38 CV windows, range 4.4%-11.0%"),
    ("Kotak Mahindra Bank", "Prophet", "7.8%", "38 CV windows, range 5.0%-10.3%"),
    ("Bank of India", "Prophet", "9.6%", "38 CV windows, range 3.8%-16.3%"),
    ("Paytm Payments Bank", "Prophet", "21.4%", "33 CV windows, high volatility"),
    ("Canara Bank", "Prophet", "N/A*", "Post-merger (Apr 2020), insufficient history for CV"),
    ("Union Bank of India", "Prophet", "N/A*", "Post-merger (Apr 2020), insufficient history for CV"),
    ("Punjab National Bank", "Prophet", "N/A*", "Post-merger (Apr 2020), insufficient history for CV"),
    ("Indian Bank", "Prophet", "N/A*", "Post-merger (Apr 2020), insufficient history for CV"),
]
add_table(["Bank", "Model", "CV MAPE (Median)", "Notes"], dc_rows)
doc.add_paragraph("* Four post-merger DC banks have <72 months of training data; CV is not yet possible. These banks will accumulate sufficient history by mid-2026.")
p = doc.add_paragraph()
run = p.add_run("DC Median (validated banks): 5.3%")
run.bold = True
doc.add_paragraph()

# ── 4. OOS ──
heading("4. Aggregate Out-of-Sample Holdout (Jan - Jun 2025)")
doc.add_paragraph("True OOS test: models trained through Dec 2024, forecasts compared against actual Jan-Jun 2025 data. This measures the combined accuracy of all individual bank models + residual after aggregation.")
oos_rows = [
    ("CC Outstanding (India)", "1.58%", "2.41%", "83% (5/6 months)"),
    ("DC Outstanding (India)", "1.02%", "2.06%", "100% (6/6 months)"),
]
add_table(["Metric", "OOS MAPE", "Max Monthly Error", "90% CI Coverage"], oos_rows)

doc.add_paragraph()
heading("5. OOS Month-by-Month Detail", level=2)
cc_oos = [
    ("Jan 2025", "10,887 Cr", "10,948 Cr", "0.6%", "Yes"),
    ("Feb 2025", "10,932 Cr", "11,043 Cr", "1.0%", "Yes"),
    ("Mar 2025", "10,989 Cr", "11,189 Cr", "1.8%", "Yes"),
    ("Apr 2025", "11,044 Cr", "11,245 Cr", "1.8%", "Yes"),
    ("May 2025", "11,105 Cr", "11,312 Cr", "1.9%", "Yes"),
    ("Jun 2025", "11,097 Cr", "11,364 Cr", "2.4%", "No"),
]
doc.add_paragraph("CC Outstanding (lakh Cr)")
add_table(["Month", "Actual", "Forecast", "APE", "In 90% CI"], cc_oos)

doc.add_paragraph()
dc_oos = [
    ("Jan 2025", "98,204 L", "1,00,182 L", "2.0%", "Yes"),
    ("Feb 2025", "98,568 L", "1,00,594 L", "2.1%", "Yes"),
    ("Mar 2025", "99,081 L", "99,810 L", "0.7%", "Yes"),
    ("Apr 2025", "99,598 L", "1,00,278 L", "0.7%", "Yes"),
    ("May 2025", "1,00,037 L", "1,00,409 L", "0.4%", "Yes"),
    ("Jun 2025", "1,00,518 L", "1,00,756 L", "0.2%", "Yes"),
]
doc.add_paragraph("DC Outstanding (lakh)")
add_table(["Month", "Actual", "Forecast", "APE", "In 90% CI"], dc_oos)
doc.add_paragraph()

# ── 6. Tiering ──
heading("6. Accuracy Tiering")
tier_rows = [
    ("Green (<=7%)", "ICICI (CC/DC), IndusInd (CC), SBI (CC/DC), HDFC (CC/DC),\nAxis (CC/DC), UCO (DC), BoB (DC), IOB (DC)"),
    ("Amber (7%-15%)", "HSBC (CC), CBI (DC), Kotak (DC), BoI (DC)"),
    ("Red (>15%)", "Kotak (CC), BoB (CC), Paytm (DC)"),
    ("Pending CV", "Yes Bank (CC), Canara (CC/DC), Union Bank (DC),\nPNB (DC), Indian Bank (DC)"),
]
add_table(["Tier", "Banks"], tier_rows)
doc.add_paragraph("Red-tier banks are in rapid growth or high-volatility phases. Dynamic logistic growth caps are computed from trailing 12-month growth to constrain over-forecasting. Pending-CV banks are generating forecasts but lack sufficient history for cross-validation.")
doc.add_paragraph()

# ── 7. Takeaways ──
heading("7. Key Takeaways")
takeaways = [
    "25 individual bank models (10 CC + 15 DC) cover 91% CC and 83% DC of India outstanding",
    "Of the 19 banks with validated CV, 15 are under 10% MAPE — well within industry-acceptable range",
    "Aggregate OOS accuracy is strong — ground-up sum achieves <2% MAPE at India level over 6 months",
    "Dual-model approach validated — ETS outperforms Prophet by 0.5-3.3pp on stable-growth banks",
    "6 banks (Yes Bank CC, Canara CC/DC, Union Bank DC, PNB DC, Indian Bank DC) are pending CV — their post-merger training windows will cross the 72-month threshold by mid-2026",
    "90% CI coverage at aggregate level is 83-100%, confirming prediction intervals are well-calibrated",
]
for i, t in enumerate(takeaways, 1):
    doc.add_paragraph(f"{i}. {t}")

out = r"D:\MPi-mip\reports\bank_forecast_accuracy_v3.docx"
doc.save(out)
print(f"Saved: {out}")
